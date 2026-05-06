"""
Offer letter generator — Claude Sonnet narrative + python-docx document.

Flow:
1. Claude Sonnet generates a professional offer letter narrative (400-500 words)
2. write_offer_docx() wraps it in a formatted .docx file
3. The file is saved to uploads/offer_letters/{application_id}.docx
4. Path stored in hr_offers.letter_path
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from tenacity import retry, stop_after_attempt, wait_exponential

logger = logging.getLogger(__name__)

CLAUDE_SONNET_MODEL = "claude-sonnet-4-6"


def build_offer_context(
    *,
    candidate_name: str,
    role: str,
    ctc_lpa: float,
    joining_date: str,
    reporting_to: str = "Admin Doshi",
    location: str = "Bangalore",
) -> dict[str, Any]:
    """Build context dict for offer letter generation."""
    return {
        "candidate_name": candidate_name,
        "role": role,
        "role_display": role.replace("_", " ").title(),
        "ctc_lpa": ctc_lpa,
        "joining_date": joining_date,
        "reporting_to": reporting_to,
        "location": location,
        "company": "YourCompany",
        "company_full": "Your Company Legal Name",
    }


class OfferGenerator:
    def __init__(self, *, client: Any) -> None:
        self._client = client

    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10), reraise=True)
    async def generate_narrative(
        self,
        *,
        candidate_name: str,
        role: str,
        ctc_lpa: float,
        joining_date: str,
        reporting_to: str = "Admin Doshi",
        location: str = "Bangalore",
    ) -> str:
        """Generate a professional offer letter narrative using Claude Sonnet, with template fallback."""
        role_display = role.replace("_", " ").title()
        try:
            prompt = (
                f"Write a professional, warm offer letter for the following hire at YourCompany "
                f"(a managed office space company in India).\n\n"
                f"Candidate: {candidate_name}\n"
                f"Role: {role_display}\n"
                f"CTC: {ctc_lpa} LPA (annual)\n"
                f"Joining Date: {joining_date}\n"
                f"Location: {location}\n"
                f"Reporting To: {reporting_to}\n\n"
                f"Write 400-500 words. Start with 'Dear {candidate_name.split()[0]},' and end with "
                f"'Warm regards, Admin Doshi, Co-Founder, YourCompany'. "
                f"Mention the role, CTC, joining date, and express genuine excitement about them joining. "
                f"Keep tone professional yet human. Do not include address blocks or legal disclaimers — "
                f"those will be added separately."
            )
            message = await self._client.messages.create(
                model=CLAUDE_SONNET_MODEL,
                max_tokens=1024,
                messages=[{"role": "user", "content": prompt}],
            )
            return message.content[0].text.strip()
        except Exception as exc:
            logger.warning("Claude offer generation failed (%s), using template fallback.", exc)
            first_name = candidate_name.split()[0]
            return (
                f"Dear {first_name},\n\n"
                f"We are delighted to extend this offer of employment to you for the position of "
                f"{role_display} at YourCompany (Your Company Legal Name).\n\n"
                f"After a thorough evaluation of your profile, experience, and the interviews conducted, "
                f"we believe you are an excellent fit for our team and are confident you will make a "
                f"significant contribution to our mission of redefining how enterprises work across India.\n\n"
                f"The details of your offer are as follows:\n\n"
                f"Position: {role_display}\n"
                f"Location: {location}\n"
                f"Date of Joining: {joining_date}\n"
                f"Reporting To: {reporting_to}\n"
                f"Annual CTC: {ctc_lpa}\n\n"
                f"We are a fast-growing managed office platform operating 500K+ sq ft of premium "
                f"workspace across India's top metros. You will be joining a driven, execution-focused "
                f"team that takes ownership and moves fast.\n\n"
                f"Please sign and return a copy of this letter to confirm your acceptance. "
                f"Should you have any questions, feel free to reach out to us at hiring@example.com.\n\n"
                f"We look forward to welcoming you to the YourCompany family.\n\n"
                f"Warm regards,\nAdmin Doshi\nCo-Founder, YourCompany"
            )


def write_offer_docx(
    *,
    output_path: str,
    candidate_name: str,
    role: str,
    ctc_lpa: float,
    joining_date: str,
    narrative: str,
    reporting_to: str = "Admin Doshi",
    location: str = "Bangalore",
) -> None:
    """Write the offer letter to a .docx file."""
    doc = Document()

    header = doc.add_heading("OFFER LETTER", level=1)
    header.alignment = 1  # Center

    doc.add_paragraph(f"Date: {joining_date}")
    doc.add_paragraph(f"To: {candidate_name}")
    doc.add_paragraph("")

    for para_text in narrative.split("\n\n"):
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph("")
    doc.add_paragraph("This offer is subject to successful background verification.")
    doc.add_paragraph("")
    doc.add_paragraph("Prepared by: Your Company Legal Name")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info("Offer letter saved: %s", output_path)


def make_offer_generator() -> OfferGenerator:
    import anthropic
    from src.config import get_settings
    client = anthropic.AsyncAnthropic(api_key=get_settings().anthropic_api_key)
    return OfferGenerator(client=client)
