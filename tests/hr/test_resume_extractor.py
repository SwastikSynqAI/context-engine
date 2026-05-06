"""Unit tests for resume text extraction."""
import io
import pytest


def test_extract_from_pdf_bytes(tmp_path):
    """Create a minimal PDF in memory and verify text extraction works."""
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        pdf_path = tmp_path / "test_resume.pdf"
        c = rl_canvas.Canvas(str(pdf_path))
        c.drawString(72, 750, "John Doe")
        c.drawString(72, 730, "5 years experience in B2B sales")
        c.save()
        pdf_bytes = pdf_path.read_bytes()
    except ImportError:
        pytest.skip("reportlab not installed — skipping PDF creation test")
        return

    from src.engines.hr.inbound.resume_extractor import extract_text_from_bytes
    text = extract_text_from_bytes(pdf_bytes, filename="test_resume.pdf")
    assert "John Doe" in text
    assert len(text) > 10


def test_extract_from_docx_bytes(tmp_path):
    """Create a minimal DOCX in memory and verify text extraction."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("Operations Manager with 8 years experience")
    docx_path = tmp_path / "test_resume.docx"
    doc.save(str(docx_path))
    docx_bytes = docx_path.read_bytes()

    from src.engines.hr.inbound.resume_extractor import extract_text_from_bytes
    text = extract_text_from_bytes(docx_bytes, filename="test_resume.docx")
    assert "Jane Smith" in text
    assert "Operations Manager" in text


def test_unsupported_extension_raises():
    from src.engines.hr.inbound.resume_extractor import extract_text_from_bytes
    with pytest.raises(ValueError, match="Unsupported"):
        extract_text_from_bytes(b"some bytes", filename="resume.txt")


def test_empty_pdf_returns_empty_string():
    """Gracefully handles a PDF that yields no text (scanned image, etc.)."""
    from src.engines.hr.inbound.resume_extractor import extract_text_from_bytes
    import io
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        buf = io.BytesIO()
        c = rl_canvas.Canvas(buf)
        c.save()
        buf.seek(0)
        pdf_bytes = buf.read()
        text = extract_text_from_bytes(pdf_bytes, filename="empty.pdf")
        assert isinstance(text, str)
    except ImportError:
        pytest.skip("reportlab not installed")


def test_detect_mime_from_filename():
    from src.engines.hr.inbound.resume_extractor import detect_file_type
    assert detect_file_type(b"", "resume.pdf") == "pdf"
    assert detect_file_type(b"", "resume.docx") == "docx"
    assert detect_file_type(b"", "RESUME.PDF") == "pdf"
